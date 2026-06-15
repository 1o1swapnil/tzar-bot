#!/usr/bin/env python3
"""Minimal Markdown -> DOCX converter tailored for the compliance analysis doc.

Handles: ATX headings (#..######), pipe tables (with header separator),
bullet lists (-/*), ordered lists, bold (**x**), inline `code`, horizontal
rules (---), blockquotes (>), and paragraphs. Not a general CommonMark engine
— just enough for our generated report. Stdlib + python-docx only.
"""
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_runs(paragraph, text):
    """Render inline **bold** and `code` spans into a paragraph."""
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        else:
            paragraph.add_run(part)


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?\s*:?-{2,}.*$", line)) and set(line.strip()) <= set("|-: ")


def parse_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def main(md_path, docx_path):
    with open(md_path, encoding="utf-8") as fh:
        lines = fh.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        # blank
        if not line.strip():
            i += 1
            continue

        # horizontal rule
        if re.match(r"^\s*---+\s*$", line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            pr = p.add_run("_" * 60)
            pr.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            i += 1
            continue

        # heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            h = doc.add_heading(level=min(level, 4))
            h.text = ""
            add_runs(h, text)
            i += 1
            continue

        # table (current line has pipes and next line is a separator)
        if "|" in line and i + 1 < n and is_table_sep(lines[i + 1]):
            header = parse_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(parse_row(lines[i]))
                i += 1
            ncol = len(header)
            table = doc.add_table(rows=1, cols=ncol)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr = table.rows[0].cells
            for c, txt in enumerate(header):
                hdr[c].paragraphs[0].text = ""
                run = hdr[c].paragraphs[0].add_run(txt)
                run.bold = True
                run.font.size = Pt(9.5)
            for row in rows:
                cells = table.add_row().cells
                for c in range(ncol):
                    txt = row[c] if c < len(row) else ""
                    para = cells[c].paragraphs[0]
                    para.text = ""
                    add_runs(para, txt)
                    for r in para.runs:
                        r.font.size = Pt(9)
            doc.add_paragraph()
            continue

        # blockquote
        if line.lstrip().startswith(">"):
            text = line.lstrip()[1:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_runs(p, text)
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ordered list
        mo = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if mo:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, mo.group(2).strip())
            i += 1
            continue

        # bullet list
        mb = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if mb:
            indent = len(mb.group(1))
            style_name = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style_name)
            add_runs(p, mb.group(2).strip())
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        add_runs(p, line.strip())
        i += 1

    doc.save(docx_path)
    print(f"wrote {docx_path}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "docs/compliance/CERT-In-CIGU-2026-0003-mapping.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else src.rsplit(".", 1)[0] + ".docx"
    main(src, dst)
